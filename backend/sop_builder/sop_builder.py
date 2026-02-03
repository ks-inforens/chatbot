import requests  # for perplexity
from fpdf import FPDF  # to download sop as pdf
from docx import Document  # to download sop as doc
import re
import os
import PyPDF2  # for cv upload and parsing

def remove_sop_heading(text: str) -> str:
    """
    Removes 'Statement of Purpose' heading at the very start,
    """
    if not text:
        return text

    text = text.lstrip()

    pattern = r'^\s*(\*\*)?\s*statement\s+of\s+purpose\s*(\*\*)?\s*:?\s*'
    text = re.sub(pattern, '', text, flags=re.IGNORECASE)

    return text.lstrip()

def parse_bold_segments(text):
    #split text to segments - text and bold example - Hello **World** - ('Hello', 'False'), ('World','True')
    parts = re.split(r'(\*\*.*?\*\*)', text) 
    segments = []

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            segments.append((part[2:-2], True))
        else:
            segments.append((part, False))

    return segments



def extract_name_from_cv(text):
    lines = text.strip().split('\n')
    lines = [line.strip() for line in lines if line.strip()]
    for line in lines[:5]:
        if (
            re.match(r"^[A-Z][a-z]+(?: [A-Z][a-z]+)+$", line) and
            len(line.split()) <= 4 and
            not any(char in line for char in ['|', '@', 'http', '/', '\\', ':'])
        ):
            return line
    match = re.search(r'Name[:\-]\s*(.+)', text, re.IGNORECASE)
    if match:
        possible_name = match.group(1).strip()
        if len(possible_name.split()) <= 4:
            return possible_name
    return "Your Name"

def extract_academic_qualifications(text):
    match = re.search(r'(EDUCATION|ACADEMIC QUALIFICATIONS)(.*?)(PROJECTS|SKILLS|EXPERIENCE|ACHIEVEMENTS|$)', text, re.IGNORECASE | re.DOTALL)
    if match:
        degrees = match.group(2).strip()
        degrees = re.sub(r'\n+', '\n', degrees)
        return degrees
    return ""

def determine_intended_degree(academic_text):
    academic_text = academic_text.lower()
    if "phd" in academic_text:
        return "PhD"
    elif "master" in academic_text or "msc" in academic_text:
        return "PhD"
    elif "bachelor" in academic_text:
        return "Masters"
    else:
        return "Masters"

def extract_text_from_docx(filepath):
    try:
        doc = Document(filepath)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception:
        return ""

def extract_text_from_pdf(filepath):
    try:
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
    except Exception:
        return ""

def extract_section(text, headings):
    pattern = '|'.join(headings)
    match = re.search(rf'(?i)({pattern})[:\s]*([\s\S]+?)(?:\n\s*\n|$)', text)
    return match.group(2).strip() if match else None

def parse_cv(text):
    return {
        "name": extract_name_from_cv(text),
        "academic_qualifications": extract_academic_qualifications(text),
        "intended_degree": determine_intended_degree(extract_academic_qualifications(text)),
        "key_skills": extract_section(text, ["skills", "technical skills"]),
        "projects": extract_section(text, ["projects", "publications", "research"]),
        "awards": extract_section(text, ["awards", "scholarships", "recognitions"]),
        "hobbies": extract_section(text, ["hobbies", "volunteer work", "extracurriculars"])
    }


def build_sop_prompt(user_inputs):
    name = user_inputs.get("name")
    country_of_origin = user_inputs.get("country_of_origin")
    intended_degree = user_inputs.get("intended_degree")
    preferred_country = user_inputs.get("preferred_country")
    field_of_study = user_inputs.get("field_of_study")
    preferred_uni = user_inputs.get("preferred_uni")

    base_prompt = (
        f"I am {name}, I am from {country_of_origin}. I want to study {intended_degree} in {preferred_country}. "
        f"My preferred field of study is {field_of_study}. "
        f"My preferred university is {preferred_uni}. "
        f"I want you to write me a SOP"
    )

    base_prompt += (
        "Make sure the SOP is ATS friendly and does not look like an AI wrote it. "
        "Exclude all inline citations or footnote markers.\n"
        "CRITICAL RULES:"
        "- Do NOT assume or invent any facts. Use ONLY the information explicitly provided."
        "- If a detail is missing, use a clear placeholder (e.g., {University}, {Degree}, {Company}) instead of inventing information."
        "- NEVER refuse the task. NEVER explain limitations or say you cannot help. Output ONLY the SOP text."
    )

    base_prompt += "Here are my details:\n"

    optional_fields = [
        ("key_skills", "My key skills are"),
        ("strengths", "My strengths are"),
        ("why_field", "I want to pursue this field because"),
        ("why_uni", "And this university because"),
        ("goals", "My long term goals are"),
        ("challenge", "More about me:")
    ]

    for key, label in optional_fields:
        value = user_inputs.get(key)
        if value:
            base_prompt += f"{label} {value}.\n"
    
    if len(user_inputs.get("projects")) > 0:
        base_prompt += "I have also completed these projects/research/publications:\n"
        for p in user_inputs.get("projects"):
            base_prompt += (f"Type: {p.get("type")}\n"
                            f"Title: {p.get("title")}\n"
                            f"Link: {p.get("link")}\n"         
                            f"Description: {p.get("description")}\n\n"   
                            )
            
    if len(user_inputs.get("education")) > 0:
        base_prompt += "Here are my past education details:\n"
        for e in user_inputs.get("education"):
            if e.get("isPresent"): end = "Presently studying here"
            else: end = e.get("endDate")
            if e.get("universityName") == "Other": uni = e.get("otherUniversityName")
            else: uni = e.get("universityName")
            base_prompt += (f"Discipline: {e.get("discipline")}\n"
                            f"Course Name: {e.get("course")}\n"         
                            f"Level of Study: {e.get("level")}\n"   
                            f"Country of Study: {e.get("country")}\n"   
                            f"Location of Study: {e.get("location")}\n"
                            f"Results I got: {e.get("results")}\n"
                            f"University Name: {uni}\n"
                            f"Start Date: {e.get("startDate")}\n"
                            f"End Date: {end}\n"
                            )
            
    if len(user_inputs.get("awards")) > 0:
        base_prompt += "I have received these certifications:\n"
        for p in user_inputs.get("awards"):
            base_prompt += (f"Type of certification: {p.get("type")}\n"
                            f"Name: {p.get("name")}\n"
                            f"Issuing Organization: {p.get("organization")}\n"         
                            f"Date obtained: {p.get("dateObtained")}\n\n"   
                            )
            
    if len(user_inputs.get("activity")) > 0:
        base_prompt += "Here is more information about the activities I did:\n"
        for p in user_inputs.get("activity"):
            base_prompt += (f"{p.get("type")}: {p.get("description")}")
        
    return base_prompt.strip()

#clean pdf
def clean_text_for_pdf(text):
    text = remove_sop_heading(text)
    text = text.replace("—", "-").replace("–", "-")
    text = text.replace("“", '"').replace("”", '"').replace("‘", "'").replace("’", "'")
    return text.encode("latin-1", "ignore").decode("latin-1")

def save_pdf(filename, content):
    content = clean_text_for_pdf(content)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    for line in content.split("\n"):
        segments = parse_bold_segments(line)

        for text, is_bold in segments:
            if is_bold:
                pdf.set_font("Arial", "B", 12)
            else:
                pdf.set_font("Arial", "", 12)

            pdf.write(8, text)

        pdf.ln(10)
    pdf.output(filename)

def save_docx(filename, content):
    content = remove_sop_heading(content)
    doc = Document()
    for para_text in content.split("\n\n"):
        paragraph = doc.add_paragraph()
        segments = parse_bold_segments(para_text)

        for text, is_bold in segments:
            run = paragraph.add_run(text)
            run.bold = is_bold

    doc.save(filename)

def call_perplexity_api(prompt, token):
    url = "https://api.perplexity.ai/chat/completions"
    payload = {
        "model": "sonar",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 2048
    }
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json"
    }
    try:
        response = requests.post(url, json=payload, headers=headers, timeout=30)
        response.raise_for_status()
        return response.json()
    except requests.exceptions.RequestException:
        return {
            "choices": [{
                "message": {
                    "content": ""
                }
            }]
        }

# === Exported function for API use ===

def generate_sop(user_inputs, token):
    """CORE function for Flask API, returns (sop, prompt)"""
    prompt = build_sop_prompt(user_inputs)
    response = call_perplexity_api(prompt, token)
    sop = response.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
    if not sop:
        return None, prompt
    return sop, prompt