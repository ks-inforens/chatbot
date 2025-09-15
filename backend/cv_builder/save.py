import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_bottom_border(paragraph):
    """Add a horizontal line under section headers"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")  # thicker line
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)

def normalize_text(text):
    if text.strip().startswith("```"):
        text = text.strip().split("```")[1]
        if text.lower().startswith("json"):
            text = text[4:].strip()
    return text.strip()


def save_as_docx(text, filename="generated_cv.docx"):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.5)    
    section.bottom_margin = Inches(0.5)  
    section.left_margin = Inches(0.5) 
    section.right_margin = Inches(0.5) 

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    try:
        data = json.loads(normalize_text(text))
        is_json = True
    except:
        is_json = False

    if is_json:
        # === Name ===
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.add_run(data.get("full_name",""))
        run.bold = True
        run.font.size = Pt(16)

        # === Location ===
        location = data.get("contact", {}).get("location", "")
        if location:
            para = doc.add_paragraph()
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = para.add_run(location)
            run.font.size = Pt(10)

        # === Contact info ===
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        contact = data.get("contact", {})
        contact_items = []
        if contact.get("email"):
            contact_items.append(contact['email'])
        if contact.get("phone"):
            contact_items.append(contact['phone'])
        if contact.get("linkedin"):
            contact_items.append(contact['linkedin'])
        para.add_run(" | ".join(contact_items))

        # === Sections ===
        sections = data.get("sections", {})
        for section, content in sections.items():
            if not content:
                continue

            # Section header
            para = doc.add_paragraph()
            run = para.add_run(section)
            run.bold = True
            run.font.size = Pt(14)
            add_bottom_border(para)
            para.paragraph_format.space_before = Pt(14)
            para.paragraph_format.space_after = Pt(8)

            # Professional Statement
            if section == "Professional Statement" and isinstance(content,str):
                doc.add_paragraph(content.strip())

            # Work Experience
            elif section == "Work Experience":
                for job in content:
                    para = doc.add_paragraph()
                    run = para.add_run(job.get("title",""))
                    run.bold = True
                    run.italic = True
                    company = job.get("company","")
                    if company:
                        run_company = para.add_run(f" | {company}")
                        run_company.bold = True
                        run_company.italic = True
                    dates = job.get("dates","")
                    if dates:
                        tab_stop = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
                        para.paragraph_format.tab_stops.add_tab_stop(tab_stop, alignment=WD_TAB_ALIGNMENT.RIGHT)
                        run_date = para.add_run(f"\t{dates}")
                        run_date.italic = True
                    for r in job.get("responsibilities", []):
                        bullet = doc.add_paragraph(r, style="List Bullet")
                        bullet.paragraph_format.space_after = Pt(2)

            # Education
            elif section == "Education":
                for edu in content:
                    para = doc.add_paragraph()
                    run = para.add_run(edu.get("institution",""))
                    run.bold = True
                    run.italic = True
                    dates = edu.get("dates","")
                    if dates:
                        tab_stop = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
                        para.paragraph_format.tab_stops.add_tab_stop(tab_stop, alignment=WD_TAB_ALIGNMENT.RIGHT)
                        run_date = para.add_run(f"\t{dates}")
                        run_date.italic = True
                    degree_field = f"{edu.get('degree','')} {edu.get('field','')}".strip()
                    if degree_field:
                        doc.add_paragraph(degree_field)
                    if edu.get("result"):
                        doc.add_paragraph(f"Result: {edu['result']}")

            # Skills / Languages (comma separated)
            elif section in ["Skills", "Languages"] and isinstance(content,list):
                doc.add_paragraph(", ".join(content))

            # Projects, Certificates, Positions, Achievements, or any new section
            else:
                if isinstance(content, list):
                    for item in content:
                        if isinstance(item, dict):
                            title = item.get("title") or item.get("position") or ""
                            company = item.get("company") or item.get("organization") or ""
                            dates = item.get("dates") or item.get("duration") or ""
                            description = item.get("description", "") or item.get("desc", "")
                            responsibilities = item.get("responsibilities",[])
                            technologies = item.get("technologies",[])

                            p = doc.add_paragraph()
                            run_main = p.add_run(title)
                            run_main.bold = True
                            run_main.italic = True
                            if company:
                                run_comp = p.add_run(f" | {company}")
                                run_comp.bold = True
                                run_comp.italic = True
                            if dates:
                                tab_stop = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
                                p.paragraph_format.tab_stops.add_tab_stop(tab_stop, alignment=WD_TAB_ALIGNMENT.RIGHT)
                                run_date = p.add_run(f"\t{dates}")
                                run_date.italic = True

                            if description:
                                doc.add_paragraph(description, style='List Bullet')
                            for r in responsibilities:
                                doc.add_paragraph(r, style='List Bullet')
                            if technologies:
                                doc.add_paragraph("Technologies: " + ", ".join(technologies))
                        elif isinstance(item,str):
                            doc.add_paragraph(item, style='List Bullet')
                elif isinstance(content,str):
                    doc.add_paragraph(content)

    else:
        # Plain text (like cover letter)
        for line in text.split("\n"):
            para = doc.add_paragraph(line.strip())
            para.paragraph_format.space_after = Pt(4)

    doc.save(filename)
    print(f"DOCX saved as {filename}")
    return filename